"""Convert all .docx files in law_data_docx/ to markdown in raw_data/."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def runs_to_md(paragraph) -> str:
    """Convert paragraph runs to markdown inline formatting, merging adjacent same-style runs."""
    # Collect (bold, italic, text) per run
    segments: list[tuple[bool, bool, str]] = []
    for run in paragraph.runs:
        if not run.text:
            continue
        segments.append((bool(run.bold), bool(run.italic), run.text))

    # Merge consecutive segments with identical formatting
    merged: list[list] = []
    for bold, italic, text in segments:
        if merged and merged[-1][0] == bold and merged[-1][1] == italic:
            merged[-1][2] += text
        else:
            merged.append([bold, italic, text])

    parts = []
    for bold, italic, text in merged:
        if bold and italic:
            text = f"***{text}***"
        elif bold:
            text = f"**{text}**"
        elif italic:
            text = f"*{text}*"
        parts.append(text)
    return "".join(parts)


def paragraph_to_md(paragraph) -> str:
    """Convert a single paragraph to a markdown line."""
    style = paragraph.style.name if paragraph.style else ""
    inline = runs_to_md(paragraph)

    # Headings
    if style.startswith("Heading 1"):
        return f"# {inline}"
    if style.startswith("Heading 2"):
        return f"## {inline}"
    if style.startswith("Heading 3"):
        return f"### {inline}"
    if style.startswith("Heading 4"):
        return f"#### {inline}"
    if style.startswith("Heading 5"):
        return f"##### {inline}"

    # List items
    num_pr = paragraph._p.find(qn("w:numPr"))
    if num_pr is not None:
        ilvl = num_pr.find(qn("w:ilvl"))
        level = int(ilvl.get(qn("w:val"), 0)) if ilvl is not None else 0
        indent = "  " * level
        return f"{indent}- {inline}"

    # Plain paragraph
    return inline


def table_to_md(table) -> list[str]:
    """Convert a table to markdown."""
    lines = []
    rows = table.rows
    if not rows:
        return lines
    for i, row in enumerate(rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
        if i == 0:
            lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return lines


def docx_to_markdown(docx_path: Path) -> str:
    """Convert a .docx file to a markdown string."""
    doc = Document(docx_path)
    lines: list[str] = []

    # Iterate over body elements preserving order (paragraphs + tables)
    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

        if tag == "p":
            # Find matching paragraph object
            para_text = None
            for p in doc.paragraphs:
                if p._p is block:
                    para_text = paragraph_to_md(p)
                    break
            if para_text is not None:
                lines.append(para_text)

        elif tag == "tbl":
            for tbl in doc.tables:
                if tbl._tbl is block:
                    lines.extend(table_to_md(tbl))
                    lines.append("")
                    break

    # Collapse consecutive blank lines
    result_lines: list[str] = []
    prev_blank = False
    for line in lines:
        is_blank = line.strip() == ""
        if is_blank and prev_blank:
            continue
        result_lines.append(line)
        prev_blank = is_blank

    return "\n".join(result_lines).strip() + "\n"


def main() -> None:
    base = Path(__file__).parent
    src_dir = base / "law_data_docx"
    dst_dir = base / "raw_data"
    dst_dir.mkdir(exist_ok=True)

    docx_files = sorted(src_dir.glob("*.docx"))
    if not docx_files:
        print("No .docx files found in law_data_docx/")
        return

    for docx_path in docx_files:
        out_path = dst_dir / (docx_path.stem + ".md")
        print(f"Converting: {docx_path.name} -> {out_path.name}")
        try:
            md = docx_to_markdown(docx_path)
            out_path.write_text(md, encoding="utf-8")
            print(f"  OK ({len(md)} chars)")
        except Exception as exc:
            print(f"  ERROR: {exc}")

    print("\nDone.")


if __name__ == "__main__":
    main()
